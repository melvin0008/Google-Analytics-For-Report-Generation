import googleanalytics as ga
import json
from docx import Document
from docx.shared import Inches
import Tkinter
import ttkcalendar
import datetime
import tkSimpleDialog
import tkMessageBox


class GoogleAnalyticsReport:
    def __init__(self,startDate,endDate,heading):
        credentials = json.load(open('secret/credentials.json'))
        accounts = ga.authenticate(**credentials)
        self.profile = accounts[0].webproperties[0].profile
        self.document = Document('secret/default.docx')
        self.document.add_heading(heading, 0)
        self.startDate = startDate
        self.endDate = endDate
        self.sessions = 0


    """
    Basic Metrics
    """
    def getBasicMetrics(self):
        self.document.add_heading('Basic Metrics', level=1)
        query = self.profile.core.query.set(metrics=['ga:users','ga:sessions','ga:pageviews','ga:uniquePageviews','ga:pageviewsPerSession','ga:bounceRate','ga:percentNewSessions']).set('start_date', self.startDate).set({'end_date': self.endDate})
        rows = query.get().rows[0]
        users = rows[0]
        self.sessions = rows[1]
        pageviews = rows[2]
        uniquePageviews = rows[3]
        pageviewsPerSession = rows[4]
        bounceRate = rows[5]
        paragraph = self.document.add_paragraph("Users: "+str(users)+"\n"+"Sessions: "+str(self.sessions)+"\n"+"Pageviews:"+ str(pageviews)+"\n"+"UniquePageviews: "+str(uniquePageviews)+"\n"+"PageviewsPerSession: "+str(pageviewsPerSession)+"\n"+"BounceRate: "+ str(bounceRate))
        
        """
        New and Returning Sessions
        """
        self.document.add_heading('New and Returning Sessions', level=1)
        percentNewSessions = rows[6]
        percentReturningSessions = 100-rows[6]
        paragraph1 = self.document.add_paragraph("PercentNewSessions: "+ str(percentNewSessions)+"\n"+"PercentReturningSessions: "+str(percentReturningSessions))


    """
    Function to Print Tables
    """
    def __print_table(self,noRows,noColumns,columnNames,rows):
        if len(rows)<noRows:
            noRows = len(rows)
        table = self.document.add_table(rows=noRows+1, cols=noColumns,style=self.document.styles['TableGrid'])
        hdr_cells = table.rows[0].cells
        for i in xrange(noColumns):
            hdr_cells[i].text = columnNames[i] 
        for i,item in enumerate(rows):
            for j in  xrange(noColumns):
                table.rows[i+1].cells[j].text = item[j] 

    """
    Function to get correct rows from the received query. Also adds column number and stringifies everything
    """
    def __get_correct_rows(self,query):
        values = query.get().rows 
        corrected_values=[]
        for i,row in enumerate(values):
            corrected_values_each=[str(i+1)+"."]
            for row_val in row:
                corrected_values_each.append(str(row_val))
            corrected_values.append(corrected_values_each) 
        return corrected_values

    """
    Top 10 pages
    """
    def getTopTenPages(self):
        self.document.add_heading("Top 10 pages", level=1)
        query = self.profile.core.query.metrics('ga:pageviews', 'ga:uniquePageviews').dimensions('ga:pathTitle').set('start_date', self.startDate).set({'end_date': self.endDate}).sort('ga:pageviews', descending=True).limit(10)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(10,4,["No.","Page Path","Pageviews","UniquePageviews"],corrected_values)

    """
    Technology used to view the website
    """
    def getTechUsedViewWebsite(self):    
        self.document.add_heading("Technology used to view the website", level=1)
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:deviceCategory').set('start_date', self.startDate).set({'end_date': self.endDate}).sort('ga:sessions', descending=True).limit(10)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(3,3,["No.","Device","Session"],corrected_values)


    """
    Audience
    """
    def getAudience(self):
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:country').set('start_date', self.startDate).set({'end_date': self.endDate}).sort('ga:sessions', descending=True)
        rows = query.get().rows
        numberOfCountries = len(rows)
        self.document.add_heading("Audience", level=1)
        paragraph1 = self.document.add_paragraph("Total Number of Countries : "+str(numberOfCountries))
        corrected_values=[]
        for i, country in enumerate(rows[:10]):
            corrected_values.append([str(i+1)+".", country[0],("%.2f" % (float(country[1]*100)/self.sessions))+"%"]) 
        self.__print_table(10,3,["No.","Country","% of Sessions"],corrected_values)


    """
    Traffic Sources
    """
    def getTrafficSources(self):
        self.document.add_heading("Traffic Sources", level=1)
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:medium').set('start_date', self.startDate).set({'end_date': self.endDate}).sort('ga:sessions', descending=True).limit(10)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(10,3,["No.","Source","Sessions"],corrected_values)

    """
    Search Terms
    """
    def getSearchTerms(self):
        self.document.add_heading("Search Terms", level=1)
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:keyword').set('start_date', self.startDate).set({'end_date': self.endDate}).sort('ga:sessions', descending=True).limit(10)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(10,3,["No.","Search Term","Sessions"],corrected_values)


    """
    Organic Sources
    """
    def getOrganicSources(self):
        self.document.add_heading("Organic Sources", level=1)
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:source').filter(medium='organic').set('start_date', self.startDate).set({'end_date': self.endDate}).limit(10).sort('ga:sessions', descending=True)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(10,3,["No.","Organic Sources","Sessions"],corrected_values)

    """
    Referral Sources
    """
    def getReferralSources(self):
        self.document.add_heading("Referral Sources", level=1)
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:source').filter(medium='referral').set('start_date', self.startDate).set({'end_date': self.endDate}).limit(10).sort('ga:sessions', descending=True)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(10,3,["No.","Referral Source","Sessions"],corrected_values)

    """
    Non Buffalo Referrals
    """
    def getNonBuffaloReferrals(self):
        self.document.add_heading("Non Buffalo Referral Sources", level=1)
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:source').filter(medium='referral').filter(source__ncontains="buffalo").set('start_date', self.startDate).set({'end_date': self.endDate}).limit(10).sort('ga:sessions', descending=True)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(10,3,["No.","Source","Sessions"],corrected_values)

    """
    Top Pages Visited as a Result of Direct Traffic
    """
    def getTopPagesFromDirectTraffic(self):
        self.document.add_heading("Top Pages Visited as a Result of Direct Traffic", level=1)
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:pathTitle').filter(medium='(none)').set('start_date', self.startDate).set({'end_date': self.endDate}).sort('ga:sessions', descending=True).limit(10)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(10,3,["No.","Pages","Sessions"],corrected_values)


    """
    Social Network
    """
    def getSocialNetwork(self):
        self.document.add_heading("Social Network", level=1)
        query = self.profile.core.query.metrics('ga:sessions').dimensions('ga:socialNetwork').set('start_date', self.startDate).set({'end_date': self.endDate}).sort('ga:sessions', descending=True).limit(10)
        corrected_values = self.__get_correct_rows(query)
        self.__print_table(10,3,["No.","Social Networks","Sessions"],corrected_values)


    def buildDoc(self):
        self.getBasicMetrics()
        self.getTopTenPages()
        self.getTechUsedViewWebsite()
        self.getAudience()
        self.getTrafficSources()
        self.getSearchTerms()
        self.getOrganicSources()
        self.getReferralSources()
        self.getNonBuffaloReferrals()
        self.getTopPagesFromDirectTraffic()
        self.getSocialNetwork()
        self.document.save('demo.docx')


class CalendarDialog(tkSimpleDialog.Dialog):
    """Dialog box that displays a calendar and returns the selected date"""
    def body(self, master):
        self.calendar = ttkcalendar.Calendar(master)
        self.calendar.pack()

    def apply(self):
        self.result = self.calendar.selection

def center(toplevel):
    toplevel.update_idletasks()
    w = toplevel.winfo_screenwidth()
    h = toplevel.winfo_screenheight()
    size = tuple(int(_) for _ in toplevel.geometry().split('+')[0].split('x'))
    x = w/2 - size[0]/2
    y = h/2 - size[1]/2
    size = (250,250)
    toplevel.geometry("%dx%d+%d+%d" % (size + (x, y)))

# Demo code:
def main():
    root = Tkinter.Tk()
    root.wm_title("UB Google analytics")
    root['bg']="#D3D8E8"
    center(root)
    dates={}

    def getStartDate():
        cd = CalendarDialog(root)
        startDateButton['text'] = cd.result.strftime("%A %d. %B %Y")
        dates['start'] = cd.result.strftime("%Y-%m-%d")

    def getEndDate():
        cd = CalendarDialog(root) 
        endDateButton['text'] = cd.result.strftime("%A %d. %B %Y")
        dates['end'] = cd.result.strftime("%Y-%m-%d")
    
    def buildDocument():
        if(startDateButton['text']!='Select Start Date!' and endDateButton['text'] != 'Select End Date!'):
            try:
                report = GoogleAnalyticsReport(dates['start'],dates['end'],'Monthly Report from '+str(startDateButton['text'])+"to "+endDateButton['text'])
                report.buildDoc()
                tkMessageBox.showinfo(title='Report Completed',message='Report Completed. Please rename your file')
            except Exception:
                 tkMessageBox.showerror(title='Not Available',message='Internet Connection needs to be checked. Also check your dates')
        else:
            tkMessageBox.showerror(title='Not Available',message='Check Your Dates')

    
    startDateButton = Tkinter.Button(root,height=2,width=20, text="Select Start Date!", command=getStartDate,bg='#3b5998',fg='white')
    startDateButton.pack(pady=10)

    endDateButton = Tkinter.Button(root,height=2,width=20, text="Select End Date!", command=getEndDate,bg='#3b5998',fg='white')
    endDateButton.pack(pady=10)

    buildButton = Tkinter.Button(root,height=3,width=20, text="Build the document", command=buildDocument,bg='#3b5998',fg='white')
    buildButton.pack(pady=10)
    root.update()
    root.mainloop()


if __name__ == "__main__":
    main() 
